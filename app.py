import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor  # 必須加上 RGBColor
from docx.oxml.ns import qn
import io
import re
from collections import Counter

st.set_page_config(page_title="變壓器節能分析系統", layout="wide")
st.title("📑 變壓器自動化分析報告 (數量精確版)")

# --- 1. 側邊欄參數設定 ---
st.sidebar.header("⚙️ 參數設定")
base_year = st.sidebar.number_input("請輸入基準年份：", value=2026)
pf_after_input = st.sidebar.number_input("設定【改善後】目標功率因數 (%)：", value=95)
age_filter = st.sidebar.selectbox("選擇變壓器齡篩選：", ["顯示全部", "超過 10 年", "超過 15 年", "超過 20 年"])

# --- 通用工具函數 ---
def set_font_kai(run, size=11, is_bold=False):
    run.font.name = '標楷體'
    run.font.size = Pt(size)
    run.font.bold = is_bold
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

def extract_number(text):
    if pd.isna(text): return 0.0
    s = str(text).replace(',', '').replace(' ', '').strip()
    nums = re.findall(r"[-+]?\d*\.\d+|\d+", s)
    return float(nums[0]) if nums else 0.0

excel_file = st.file_uploader("請上傳您的 Excel 檔案", type=["xlsx"])

if excel_file:
    # 讀取 Excel
    raw_df = pd.read_excel(excel_file, sheet_name=0, header=None)
    all_transformer_data = []
    seen_sn = set() # 儲存已抓取的變壓器編號

    # 定位「序號」座標
    anchors = []
    for r in range(len(raw_df)):
        for c in range(len(raw_df.columns)):
            if str(raw_df.iloc[r, c]).replace(' ', '') == "序號":
                # 檢查下方是否為規格表的特徵標籤
                if r + 1 < len(raw_df):
                    next_v = str(raw_df.iloc[r+1, c]).replace(' ', '')
                    if any(k in next_v for k in ["建築", "編號", "位置"]):
                        anchors.append((r, c))
    
    # 遍歷所有找到的規格表區塊
    for r_start, c_start in anchors:
        # 橫向掃描設備 (精確掃描 TR-1 到 TR-7)
        for offset in range(1, 10): 
            target_col = c_start + offset
            if target_col >= len(raw_df.columns): break
            
            # 初始化設備字典
            d = {"建築物": "-", "編號": "-", "年份": 0, "廠牌": "-", "容量": 0.0, 
                 "型式": "-", "負載率": 0.0, "現況功因": 0.0, "鐵損": 0.0, "實際銅損": 0.0, "改善前耗能": 0.0}
            specs = []
            
            # 垂直掃描該欄位的內容
            is_valid_device = False
            for r_offset in range(0, 45):
                curr_r = r_start + r_offset
                if curr_r >= len(raw_df): break
                
                # 標籤偵測
                l1 = str(raw_df.iloc[curr_r, c_start]).strip()
                l2 = str(raw_df.iloc[curr_r, c_start-1]).strip() if c_start > 0 else ""
                label = l1 if (l1 != "nan" and l1 != "") else l2
                lp = label.replace(' ', '').replace('\n', '')
                
                if r_offset > 0 and lp == "序號": break
                
                val = str(raw_df.iloc[curr_r, target_col]).strip()
                if lp == "nan" or not lp or val == "nan": continue

                # 資料分類抓取
                if any(k in lp for k in ["建築", "位置"]): d["建築物"] = val
                if "編號" in lp: 
                    d["編號"] = val
                    is_valid_device = True # 只要有編號就視為潛在設備
                if "廠牌" in lp: d["廠牌"] = val
                if "型式" in lp: d["型式"] = val
                if any(k in lp for k in ["年份", "出廠"]):
                    n = extract_number(val)
                    d["年份"] = int(n) + 1911 if 0 < n < 200 else int(n)
                if "容量" in lp: d["容量"] = extract_number(val)
                if any(k in lp for k in ["利用率", "負載率"]):
                    n = extract_number(val)
                    d["負載率"] = n * 100 if 0 < n < 1 else n
                if lp == "功因" or any(k in lp for k in ["功率因數", "PF", "P.F"]):
                    n_pf = extract_number(val)
                    if n_pf > 0: d["現況功因"] = n_pf / 100 if n_pf > 1 else n_pf

                specs.append((label, val))

            # 最終儲存判斷：必須有容量、有編號，且沒重複抓過
            if is_valid_device and d["容量"] > 0:
                # 建立唯一 Key (建築物+編號)
                unique_key = f"{d['建築物']}_{d['編號']}"
                if unique_key not in seen_sn:
                    # 功因預設
                    if d["現況功因"] <= 0: d["現況功因"] = 0.8
                    
                    # 計算公式
                    d["鐵損"] = d["容量"] * 3.5
                    d["實際銅損"] = (d["容量"] * 13.0) * ((d["負載率"] / 100) ** 2)
                    d["改善前耗能"] = (d["鐵損"] + d["實際銅損"]) * 8760 / 1000
                    
                    # 篩選邏輯
                    age = base_year - d["年份"] if d["年份"] > 0 else 0
                    if (age_filter == "超過 10 年" and age < 10) or \
                       (age_filter == "超過 15 年" and age < 15) or \
                       (age_filter == "超過 20 年" and age < 20): continue
                    
                    all_transformer_data.append({"specs": specs, "analysis": d})
                    seen_sn.add(unique_key)

      if all_transformer_data:
        # --- 數據摘要顯示 ---
        total_cap = sum(t["analysis"]["容量"] for t in all_transformer_data)
        cap_counts = Counter(t["analysis"]["容量"] for t in all_transformer_data)
        dist_str = "、".join([f"{k}kVA x {v}台" for k, v in sorted(cap_counts.items(), reverse=True)])
        # --- 關鍵修正：務必在顯示 st.metric 之前加入這一行 ---
        avg_usage = sum(t["analysis"]["負載率"] for t in all_transformer_data) / len(all_transformer_data)

        st.success(f"✅ 解析完成！符合篩選條件：共 {len(all_transformer_data)} 台")
        c1, c2 = st.columns(2)
        with c1:
            st.metric("1. 總裝置容量", f"{total_cap:,.0f} kVA")
            st.write("**2. 規格台數分布：**")
            for k, v in sorted(cap_counts.items(), reverse=True):
                st.write(f"　🔹 {k:,.0f} kVA × {v} 台")
        with c2:
            st.metric("3. 平均負載利用率", f"{avg_usage:.2f} %")

        # --- Word 產出邏輯 ---
        doc = Document()
        # 壹、總表
        doc.add_heading('壹、 設備統計總表', 1)
        st_t = doc.add_table(rows=0, cols=2); st_t.style = 'Table Grid'
        def add_sum(l, v):
            r = st_t.add_row().cells
            set_font_kai(r[0].paragraphs[0].add_run(l), 12, True)
            set_font_kai(r[1].paragraphs[0].add_run(v), 12)
        add_sum("總裝置容量", f"{total_cap:,.0f} kVA")
        add_sum("設備規格分布", "、".join([f"{k}kVA x {v}台" for k, v in sorted(cap_counts.items(), reverse=True)]))
        add_sum("平均負載利用率", f"{avg_usage:.2f} %")
        doc.add_page_break()

        # 貳、數據分析表
        doc.add_heading('貳、 變壓器設備改善前數據分析表', 1)
        ana_t = doc.add_table(rows=1, cols=11); ana_t.style = 'Table Grid'
        headers = ["建築物", "編號", "年份", "廠牌", "容量", "型式", "負載率", "現況功因", "銅損(W)", "鐵損(W)", "改善前耗能"]
        for i, h in enumerate(headers): set_font_kai(ana_t.rows[0].cells[i].paragraphs[0].add_run(h), 8, True)
        for t in all_transformer_data:
            d = t["analysis"]
            row = ana_t.add_row().cells
            row_vals = [d["建築物"], d["編號"], d["年份"], d["廠牌"], f"{d['容量']:.0f}", d["型式"], f"{d['負載率']:.1f}%", f"{d['現況功因']:.2f}", f"{d['實際銅損']:.1f}", f"{d['鐵損']:.1f}", f"{int(d['改善前耗能']):,}"]
            for i, v in enumerate(row_vals): set_font_kai(row[i].paragraphs[0].add_run(str(v)), 8)
        
        doc.add_page_break()
        # 參、詳細資料
        doc.add_heading('參、 詳細設備數據', 1)
        for t in all_transformer_data:
            doc.add_paragraph().add_run(f"設備詳細資料 (編號：{t['analysis']['編號']})").bold = True
            dt = doc.add_table(rows=0, cols=2); dt.style = 'Table Grid'
            for l, v in t["specs"]:
                row = dt.add_row().cells
                set_font_kai(row[0].paragraphs[0].add_run(l), 10); set_font_kai(row[1].paragraphs[0].add_run(v), 10)
            doc.add_page_break()
# ==========================================================
        # 肆、 節能改善建議報告 (制式文字與紅字帶入)
        # ==========================================================
        doc.add_page_break()
        doc.add_heading('肆、 節能改善建議報告', 1)

        # --- 計算報告所需動態數值 ---
        # 1. 預估改善後耗能 (參考 AMT 表格，假設鐵損降至 1/5, 銅損降至 0.6倍)
        # 你也可以根據之前的 IRON_LOSS_MAP 另外建立一個 AMT_LOSS_MAP 做查表
        total_kwh_before = sum(t["analysis"]["改善前耗能"] for t in all_transformer_data)
        total_kwh_after = total_kwh_before * 0.35  # 估計節電率 65%
        savings_kwh = total_kwh_before - total_kwh_after
        savings_money = savings_kwh * 3.3          # 假設電費 3.3 元/度
        
        # 2. 投資費用預估 (這部分建議依你的實際單價調整)
        invest_cost = total_cap * 1600        # 假設每 kVA 成本 1600 元
        payback_year = (invest_cost / savings_money) if savings_money > 0 else 0

        # --- 一、 現況說明 ---
        doc.add_heading('一、 現況說明', 2)
        p1 = doc.add_paragraph()
        p1.add_run("1. 依據非生產性質能源查核申報資料，貴單位高壓變壓器總裝置容量達 ")
        p1.add_run(f"{total_cap:,.0f} kVA").font.color.rgb = RGBColor(255, 0, 0) # 紅字
        p1.add_run("，平常雖然注重保養維持正常運轉，但效率與新型非晶質高效率變壓器相比，其無載損耗(kW)基本差異大。現況使用 20 年以上。")

        p2 = doc.add_paragraph()
        p2.add_run("2. 依據查核系統資料，評估 ")
        p2.add_run(f"{dist_str}").font.color.rgb = RGBColor(255, 0, 0)
        p2.add_run(" 台變壓器現況年平均利用率 ")
        p2.add_run(f"{avg_usage:.1f} %").font.color.rgb = RGBColor(255, 0, 0) # 紅字
        p2.add_run("，其變壓器計算總損失約 ")
        p2.add_run(f"{(total_kwh_before/8760):.2f} kW").font.color.rgb = RGBColor(255, 0, 0) # 紅字
        p2.add_run("，以 1 年 8760 小時運轉，推估計算年耗能約為 ")
        p2.add_run(f"{total_kwh_before:,.0f} kWh/年").font.color.rgb = RGBColor(255, 0, 0) # 紅字
        p2.add_run("。")

        # --- 二、 改善方案 ---
        doc.add_heading('二、 改善方案', 2)
        p3 = doc.add_paragraph("變壓器的損失有二種：一種發生在變壓器和配電線路接續時所產生的無負載損失（鐵損），另一種是在使用電力時才會發生的負載損失（銅損）。為了降低變壓器的無載損失（鐵損），建議將傳統的矽鋼片鐵心，改採用高性能的非晶質合金材料(Amorphous Alloy)。其鐵損是現況方向性矽鋼片的 1/3-1/5，可降低變壓器損失。")
        
        # (這裡可以插入你截圖中的對照表，或維持文字說明)
        p3_sub = doc.add_paragraph("非晶質變壓器優點：(1)鐵心結構、噪音較低 5~8dB，低損耗、低運轉溫度。(2)節能效果明顯，比一般型降低 20%~40% 以上。")

        # --- 三、 預期效益 ---
        doc.add_heading('三、 預期效益', 2)
        p4 = doc.add_paragraph()
        p4.add_run("1. 預期效益：建議可規劃將傳統鐵心式變壓器汰換為高效率非晶質變壓器，其節能效益推估計算約可減少 ")
        p4.add_run(f"{savings_kwh:,.0f} kWh/年").font.color.rgb = RGBColor(255, 0, 0) # 紅字
        p4.add_run("，節省電費 ")
        p4.add_run(f"{(savings_money/10000):.1f} 萬元/年").font.color.rgb = RGBColor(255, 0, 0) # 紅字
        p4.add_run("。")

        p5 = doc.add_paragraph()
        p5.add_run("2. 投資費用：高效率變壓器汰換投資費用預估約 ")
        p5.add_run(f"{(invest_cost/10000):.1f} 萬元").font.color.rgb = RGBColor(255, 0, 0) # 紅字
        p5.add_run(" (實際金額依廠商報價為主)。")

        p6 = doc.add_paragraph()
        p6.add_run("3. 回收年限：")
        p6.add_run(f"{(invest_cost/10000):.1f} 萬元 ÷ {(savings_money/10000):.1f} 萬元/年 = ").font.color.rgb = RGBColor(255, 0, 0) # 紅字
        p6.add_run(f"{payback_year:.1f} 年").font.color.rgb = RGBColor(255, 0, 0) # 紅字
        p6.add_run("。")
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        st.download_button("📥 下載完整報告", output, "Transformer_Report_Final.docx")
