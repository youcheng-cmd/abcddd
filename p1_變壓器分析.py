import streamlit as st
import pandas as pd
import os  # 移到這裡
from docx import Document
from docx.shared import Pt, RGBColor, Inches  # 移到這裡
from docx.oxml.ns import qn
import io
import re
from collections import Counter
# --- 1. 資料抓取邏輯 (全域優先) ---
raw_df = None 

# 優先檢查全域暫存
if 'global_excel' in st.session_state and st.session_state['global_excel'] is not None:
    try:
        uploaded_file = st.session_state['global_excel']
        xl = pd.ExcelFile(uploaded_file)
        target_sheet = None
        for s in xl.sheet_names:
            if any(k in s for k in ["表八", "電能系統資料", "表8"]):
                target_sheet = s
                break
        if target_sheet:
            raw_df = pd.read_excel(uploaded_file, sheet_name=target_sheet, header=None)
            st.info(f"💡 已自動從全域檔案抓取：{target_sheet}")
    except Exception as e:
        st.warning(f"全域檔案讀取失敗：{e}")

# 若全域沒檔案，才顯示本頁上傳框
if raw_df is None:
    uploaded_file_local = st.file_uploader("請上傳您的 Excel 檔案 (包含表八)", type=["xlsx"], key="local_p1_up")
    if uploaded_file_local:
        raw_df = pd.read_excel(uploaded_file_local, header=None)
# --- 以下接原本的 AMT_SPECS 資料庫與後續計算 ---
# --- 這裡放資料庫 ---
AMT_SPECS = {
    100: [49, 0.07, 1.15],
    150: [61, 0.08, 1.68],
    200: [73.4, 0.11, 1.96],
    300: [84, 0.13, 2.19],
    400: [96.4, 0.15, 2.40],
    500: [116, 0.20, 3.20],
    600: [135.8, 0.24, 4.00],
    750: [142, 0.28, 6.00],
    1000: [169, 0.40, 8.30],
    1250: [195.2, 0.45, 10.25],
    1500: [215.4, 0.50, 12.10],
    2000: [285.2, 0.70, 14.10],
    2500: [332, 0.80, 16.57],
    3000: [360.4, 0.90, 17.25]
}

# --- 這裡放計算邏輯函數 ---
def get_best_amt_cap(old_cap, current_load_factor):
    load_kva = old_cap * (current_load_factor / 100)
    for cap in sorted(AMT_SPECS.keys()):
        new_lf = (load_kva / cap) * 100
        if 30 <= new_lf <= 35:
            return cap
    return old_cap if old_cap in AMT_SPECS else min(AMT_SPECS.keys(), key=lambda x:abs(x-old_cap))

st.title("📑 變壓器自動化分析報告 (數量精確版)")

# --- 1. 側邊欄參數設定 ---
st.sidebar.header("⚙️ 參數設定")
base_year = st.sidebar.number_input("請輸入基準年份：", value=2026)
pf_after_input = st.sidebar.number_input("設定【改善後】目標功率因數 (%)：", value=95)
age_filter = st.sidebar.selectbox("選擇變壓器齡篩選：", ["顯示全部", "超過 10 年", "超過 15 年", "超過 20 年"])
# 將電費輸入框移至側邊欄 (sidebar)
electricity_price = st.sidebar.number_input("請輸入平均電費 (元/度)", min_value=0.0, value=5.00, step=0.01)

# --- 通用工具函數 ---
def set_font_kai(run, size=12, is_bold=False, color=RGBColor(0, 0, 0)):
    run.font.name = '標楷體'
    run.font.size = Pt(size)
    run.font.bold = is_bold
    run.font.color.rgb = color
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

def extract_number(text):
    if pd.isna(text): return 0.0
    s = str(text).replace(',', '').replace(' ', '').strip()
    nums = re.findall(r"[-+]?\d*\.\d+|\d+", s)
    return float(nums[0]) if nums else 0.0

if raw_df is not None:
    # 讀取 Excel
    all_transformer_data = []
    seen_sn = set() # 儲存已抓取的變壓器編號

    # 定位「序號」座標
    anchors = []
    for r in range(len(raw_df)):
        for c in range(len(raw_df.columns)):
            if str(raw_df.iloc[r, c]).replace(' ', '') == "序號":
                # 檢查下方是否為規格表的特徵標籤
                if r + 1 < len(raw_df):
                    next_v = str(raw_df.iloc[r+1, c]).replace(' ', '')
                    if any(k in next_v for k in ["建築", "編號", "位置"]):
                        anchors.append((r, c))
    
    # 遍歷所有找到的規格表區塊
    for r_start, c_start in anchors:
        # 橫向掃描設備 (精確掃描 TR-1 到 TR-7)
        for offset in range(1, 10): 
            target_col = c_start + offset
            if target_col >= len(raw_df.columns): break
            
            # 初始化設備字典
            d = {"建築物": "-", "編號": "-", "年份": 0, "廠牌": "-", "容量": 0.0, 
                 "型式": "-", "負載率": 0.0, "現況功因": 0.0, "鐵損": 0.0, "實際銅損": 0.0, "改善前耗能": 0.0}
            specs = []
            
            # 垂直掃描該欄位的內容
            is_valid_device = False
            for r_offset in range(0, 45):
                curr_r = r_start + r_offset
                if curr_r >= len(raw_df): break
                
                # 標籤偵測
                l1 = str(raw_df.iloc[curr_r, c_start]).strip()
                l2 = str(raw_df.iloc[curr_r, c_start-1]).strip() if c_start > 0 else ""
                label = l1 if (l1 != "nan" and l1 != "") else l2
                lp = label.replace(' ', '').replace('\n', '')
                
                if r_offset > 0 and lp == "序號": break
                
                val = str(raw_df.iloc[curr_r, target_col]).strip()
                if lp == "nan" or not lp or val == "nan": continue

                # 資料分類抓取
                if any(k in lp for k in ["建築", "位置"]): d["建築物"] = val
                if "編號" in lp: 
                    d["編號"] = val
                    is_valid_device = True # 只要有編號就視為潛在設備
                if "廠牌" in lp: d["廠牌"] = val
                if "型式" in lp: d["型式"] = val
                if any(k in lp for k in ["年份", "出廠"]):
                    n = extract_number(val)
                    d["年份"] = int(n) + 1911 if 0 < n < 200 else int(n)
                if "容量" in lp: d["容量"] = extract_number(val)
                if any(k in lp for k in ["利用率", "負載率"]):
                    n = extract_number(val)
                    d["負載率"] = n * 100 if 0 < n < 1 else n
                if lp == "功因" or any(k in lp for k in ["功率因數", "PF", "P.F"]):
                    n_pf = extract_number(val)
                    if n_pf > 0: d["現況功因"] = n_pf / 100 if n_pf > 1 else n_pf

                specs.append((label, val))

            # 最終儲存判斷：必須有容量、有編號，且沒重複抓過
            if is_valid_device and d["容量"] > 0:
                # 建立唯一 Key (建築物+編號)
                unique_key = f"{d['建築物']}_{d['編號']}"
                if unique_key not in seen_sn:
                    # 功因預設
                    if d["現況功因"] <= 0: d["現況功因"] = 0.8
                    
                    # 計算公式
                    d["鐵損"] = d["容量"] * 3.5
                    d["實際銅損"] = (d["容量"] * 13.0) * ((d["負載率"] / 100) ** 2)
                    d["改善前耗能"] = (d["鐵損"] + d["實際銅損"]) * 8760 / 1000
                    
                    # 篩選邏輯
                    age = base_year - d["年份"] if d["年份"] > 0 else 0
                    if (age_filter == "超過 10 年" and age < 10) or \
                       (age_filter == "超過 15 年" and age < 15) or \
                       (age_filter == "超過 20 年" and age < 20): continue
                    
                    all_transformer_data.append({"specs": specs, "analysis": d})
                    seen_sn.add(unique_key)

if all_transformer_data:
        # --- 數據摘要顯示 ---
        total_cap = sum(t["analysis"]["容量"] for t in all_transformer_data)
        cap_counts = Counter(t["analysis"]["容量"] for t in all_transformer_data)
        dist_str = "、".join([f"{k:,.0f}kVA x {v}台" for k, v in sorted(cap_counts.items(), reverse=True)])
        # --- 關鍵修正：務必在顯示 st.metric 之前加入這一行 ---
        avg_usage = sum(t["analysis"]["負載率"] for t in all_transformer_data) / len(all_transformer_data)

        st.success(f"✅ 解析完成！符合篩選條件：共 {len(all_transformer_data)} 台")
        c1, c2 = st.columns(2)
        with c1:
            st.metric("1. 總裝置容量", f"{total_cap:,.0f} kVA")
            st.write("**2. 規格台數分布：**")
            for k, v in sorted(cap_counts.items(), reverse=True):
                st.write(f"　🔹 {k:,.0f} kVA × {v} 台")
        with c2:
            st.metric("3. 平均負載利用率", f"{avg_usage:.2f} %")

        # --- Word 產出邏輯 ---
        doc = Document()
        # 壹、總表
        doc.add_heading('壹、 設備統計總表', 1)
        st_t = doc.add_table(rows=0, cols=2); st_t.style = 'Table Grid'
        def add_sum(l, v):
            r = st_t.add_row().cells
            set_font_kai(r[0].paragraphs[0].add_run(l), 12, True)
            set_font_kai(r[1].paragraphs[0].add_run(v), 12)
        add_sum("總裝置容量", f"{total_cap:,.0f} kVA")
        add_sum("設備規格分布", "、".join([f"{k}kVA x {v}台" for k, v in sorted(cap_counts.items(), reverse=True)]))
        add_sum("平均負載利用率", f"{avg_usage:.2f} %")
        doc.add_page_break()

        # 貳、數據分析表
        doc.add_heading('貳、 變壓器設備改善前數據分析表', 1)
        ana_t = doc.add_table(rows=1, cols=11); ana_t.style = 'Table Grid'
        headers = ["建築物", "編號", "年份", "廠牌", "容量", "型式", "負載率", "現況功因", "銅損(W)", "鐵損(W)", "改善前耗能"]
        for i, h in enumerate(headers): set_font_kai(ana_t.rows[0].cells[i].paragraphs[0].add_run(h), 8, True)
        for t in all_transformer_data:
            d = t["analysis"]
            row = ana_t.add_row().cells
            row_vals = [d["建築物"], d["編號"], d["年份"], d["廠牌"], f"{d['容量']:.0f}", d["型式"], f"{d['負載率']:.1f}%", f"{d['現況功因']:.2f}", f"{d['實際銅損']:.1f}", f"{d['鐵損']:.1f}", f"{int(d['改善前耗能']):,}"]
            for i, v in enumerate(row_vals): set_font_kai(row[i].paragraphs[0].add_run(str(v)), 8)
        
        doc.add_page_break()
        # 參、詳細資料
        doc.add_heading('參、 詳細設備數據', 1)
        for t in all_transformer_data:
            doc.add_paragraph().add_run(f"設備詳細資料 (編號：{t['analysis']['編號']})").bold = True
            dt = doc.add_table(rows=0, cols=2); dt.style = 'Table Grid'
            for l, v in t["specs"]:
                row = dt.add_row().cells
                set_font_kai(row[0].paragraphs[0].add_run(l), 10); set_font_kai(row[1].paragraphs[0].add_run(v), 10)
            doc.add_page_break()
# ==========================================================
        # 肆、 節能改善建議報告 (制式文字與紅字帶入)
        # ==========================================================
        doc.add_page_break()
        title_p = doc.add_heading('', 1)
        set_font_kai(title_p.add_run('節能改善建議報告'), size=14, is_bold=True)

        # --- 計算報告所需動態數值 ---
        # 1. 預估改善後耗能 (參考 AMT 表格，假設鐵損降至 1/5, 銅損降至 0.6倍)
        # 你也可以根據之前的 IRON_LOSS_MAP 另外建立一個 AMT_LOSS_MAP 做查表
        total_kwh_before = sum(t["analysis"]["改善前耗能"] for t in all_transformer_data)
        total_kwh_after = total_kwh_before * 0.35  # 估計節電率 65%
        savings_kwh = total_kwh_before - total_kwh_after
        savings_money = savings_kwh * 3.3          # 假設電費 3.3 元/度
        
        # 2. 投資費用預估 (這部分建議依你的實際單價調整)
        invest_cost = total_cap * 1600        # 假設每 kVA 成本 1600 元
        payback_year = (invest_cost / savings_money) if savings_money > 0 else 0

        # --- 一、 現況說明 ---
        h1 = doc.add_paragraph()
        set_font_kai(h1.add_run('一、現況說明'), size=14, is_bold=True)
        p1 = doc.add_paragraph()
        # 固定文字為黑色 12號
        set_font_kai(p1.add_run("1.依據非生產性質能源查核申報資料，貴單位高壓變壓器總裝置容量達 "), size=12)
        # 變動數據為紅色 12號
        set_font_kai(p1.add_run(f"{total_cap:,.0f} kVA"), size=12, color=RGBColor(255, 0, 0))
        set_font_kai(p1.add_run("，平常雖然注重保養維持正常運轉，但效率與新型非晶質高效率變壓器相比，其無載損耗(kW)基本差異大。現況使用 20 年以上。"), size=12)
        p2 = doc.add_paragraph()
        set_font_kai(p2.add_run("2.依據查核系統資料，評估 "), size=12)
        # 修改後的 dist_str 已不帶 .0
        set_font_kai(p2.add_run(dist_str), size=12, color=RGBColor(255, 0, 0))
        set_font_kai(p2.add_run(" 台變壓器現況年平均利用率 "), size=12)
        set_font_kai(p2.add_run(f"{avg_usage:.1f}%"), size=12, color=RGBColor(255, 0, 0))
        set_font_kai(p2.add_run("，其變壓器計算總損失約 "), size=12)
        set_font_kai(p2.add_run(f"{(total_kwh_before/8760):.2f} kW"), size=12, color=RGBColor(255, 0, 0))
        set_font_kai(p2.add_run("，以 1 年 8760 小時運轉，推估計算年耗能約為 "), size=12)
        set_font_kai(p2.add_run(f"{total_kwh_before:,.0f} kWh/年"), size=12, color=RGBColor(255, 0, 0))
        set_font_kai(p2.add_run("。"), size=12)
# --- 插入現況數據表格 (就在 P2 下面) ---
        summary_table = doc.add_table(rows=1, cols=11)
        summary_table.style = 'Table Grid'
        summary_table.alignment = 1 

        headers = ["建築物", "編號", "年份", "廠牌", "容量", "型式", "負載率", "現況功因", "銅損(W)", "鐵損(W)", "改善前耗能"]
        header_cells = summary_table.rows[0].cells
        for i, h in enumerate(headers):
            p = header_cells[i].paragraphs[0]
            p.alignment = 1
            run = p.add_run(h)
            set_font_kai(run, size=9, is_bold=True)

        for t in all_transformer_data:
            d = t["analysis"]
            row_cells = summary_table.add_row().cells
            # 依據你的截圖，數值格式精確化
            vals = [
                str(d["建築物"]), str(d["編號"]), str(d["年份"]), str(d["廠牌"]), 
                f"{d['容量']:,.0f}", str(d["型式"]), f"{d['負載率']:.1f}%", 
                f"{d['現況功因']:.2f}", f"{d['實際銅損']:,.1f}", 
                f"{d['鐵損']:,.0f}", f"{int(d['改善前耗能']):,}"
            ]
            for i, v in enumerate(vals):
                p = row_cells[i].paragraphs[0]
                p.alignment = 1
                set_font_kai(p.add_run(v), size=8)
        # --- 二、 改善方案 ---
        h2 = doc.add_paragraph()
        set_font_kai(h2.add_run('二、改善方案'), size=14, is_bold=True)

        # 1. 插入原理說明文字 (這段現在放在最前面)
        p_desc = doc.add_paragraph()
        run_desc = p_desc.add_run("1.變壓器的損失有二種：一種發生在變壓器和配電線路接續時所產生的無負載損失（鐵損），另一種是在使用電力時才會發生的負載損失（銅損）。在配電線路連續供電負載下，二種損失的比較而言，無負載損失較大，對能源的使用是一種很大的耗損。為了降低變壓器的無負載損失（鐵損），將傳統的方向性矽鋼片鐵心，改採用高性能的非晶質合金材料（Amorphous Alloy），其鐵損是現況方向性矽鋼片的 1/3-1/5，可降低變壓器損失，下表為油式非晶質合金與方向性矽鋼片的銅、鐵損失比較。")
        set_font_kai(run_desc, size=12)

        # 2. 插入表格標題 (image1 上方)
        p_table_title = doc.add_paragraph()
        p_table_title.alignment = 1 # 標題置中
        run_title = p_table_title.add_run("11.4/22.8kV 一般傳統鐵心矽鋼片與非晶質(AMT)變壓器銅、鐵損參考表")
        set_font_kai(run_title, size=12, is_bold=True)

        # 3. 插入 image1 (表格圖，設定為 4 英吋)
        if os.path.exists('image1.png'):
            doc.add_picture('image1.png', width=Inches(4))
            doc.paragraphs[-1].alignment = 1 # 圖片置中
            doc.add_paragraph() # 空行

        # 4. 插入三點特點文字
        features = [
            "2.非晶質變壓器(圖一)優點:",
            "(1) 鐵心結構，噪音較低 5~6dB，低損耗、低運轉溫度，有效延長使用設備壽命。",
            "(2) 低損耗，耗能較矽鋼片變壓器降低 20%~40% 以上。",
            "(3) 變壓器為非晶合金製作低耗能，減少 SO2、CO2 及 NOX 的排放量，可緩和溫室效應及環境保護。"
        ]
        for f in features:
            pf = doc.add_paragraph()
            run_f = pf.add_run(f)
            set_font_kai(run_f, size=12)

        doc.add_paragraph() # 預留一點間距

        # 5. 插入並排圖片 (image2 與 image3，各設定為 2.3 英吋)
        img_table = doc.add_table(rows=1, cols=2)
        img_table.alignment = 1 
        
        # 左欄放入 image2
        if os.path.exists('image2.png'):
            cell_l = img_table.rows[0].cells[0]
            para_l = cell_l.paragraphs[0]
            para_l.alignment = 1
            para_l.add_run().add_picture('image2.png', width=Inches(2.3))
            
        # 右欄放入 image3
        if os.path.exists('image3.png'):
            cell_r = img_table.rows[0].cells[1]
            para_r = cell_r.paragraphs[0]
            para_r.alignment = 1
            para_r.add_run().add_picture('image3.png', width=Inches(2.3))

        # 6. 並排圖片下方的統一圖說
        p_img_caption = doc.add_paragraph()
        p_img_caption.alignment = 1
        run_cap = p_img_caption.add_run("圖一、非晶質乾式及油浸式變壓器")
        set_font_kai(run_cap, size=12, is_bold=True)
        # --- 三、 預期效益 ---
        h3 = doc.add_paragraph()
        set_font_kai(h3.add_run('三、預期效益'), size=14, is_bold=True)
# --- 三、 預期效益 表格修正 (改為 10 欄) ---
        benefit_table = doc.add_table(rows=1, cols=10) # 欄數從 9 改成 10
        benefit_table.style = 'Table Grid'
        benefit_table.alignment = 1

        # 新增「投資(萬)」欄位
        headers = ["編號", "更換容量", "負載%", "銅損(W)", "鐵損(W)", "損耗kWh", "節電量", "改善費用(萬)", "節省(萬)", "回收年"]
        header_cells = benefit_table.rows[0].cells
        for i, h in enumerate(headers):
            p = header_cells[i].paragraphs[0]
            p.alignment = 1
            set_font_kai(p.add_run(h), size=8, is_bold=True) # 欄位變多，表頭建議縮小到 8 級

        # 初始化加總變數
        savings_kwh = 0
        savings_money = 0 
        invest_cost = 0   

        for t in all_transformer_data:
            old = t["analysis"]
            new_cap = get_best_amt_cap(old['容量'], old['負載率'])
            spec = AMT_SPECS[new_cap]
            
            # 計算改善後數值
            new_lf = (old['容量'] * (old['負載率']/100) / new_cap) * 100
            new_wcu = spec[2] * ((new_lf/100)**2) * 1000 
            new_wfe = spec[1] * 1000 
            new_kwh = (new_wcu + new_wfe) / 1000 * 8760
            
            s_kwh = old['改善前耗能'] - new_kwh
            s_money = s_kwh * electricity_price
            inv_wan = spec[0] # 單台投資金額 (萬元)
            pb = inv_wan / (s_money/10000) if s_money > 0 else 0
            
            # 填入表格數據 (共 10 項)
            row = benefit_table.add_row().cells
            row_vals = [
                old['編號'], 
                f"{new_cap:,.0f}", 
                f"{new_lf:.1f}%", 
                f"{new_wcu:,.1f}", 
                f"{new_wfe:,.1f}", 
                f"{new_kwh:,.0f}", 
                f"{s_kwh:,.0f}", 
                f"{inv_wan:.1f}",   # 新增：單台改善費用
                f"{(s_money/10000):.2f}", 
                f"{pb:.1f}"
            ]
            
            for i, v in enumerate(row_vals):
                p = row[i].paragraphs[0]
                p.alignment = 1
                set_font_kai(p.add_run(str(v)), size=8) # 數據字體建議維持 8 級

            # 累加總數
            savings_kwh += s_kwh
            savings_money += s_money
            invest_cost += (inv_wan * 10000) # 轉回「元」供下方文字計算

        # 計算總回收年限
        payback_year = (invest_cost / savings_money) if savings_money > 0 else 0
        
        doc.add_paragraph() # 表格完空一行再接 p4
        p4 = doc.add_paragraph()
        set_font_kai(p4.add_run("1. 預期效益：建議可規劃將傳統鐵心式變壓器汰換為高效率非晶質變壓器，其節能效益推估計算約可減少 "), size=12)
        set_font_kai(p4.add_run(f"{savings_kwh:,.0f} kWh/年"), size=12, color=RGBColor(255, 0, 0)) # 紅字
        set_font_kai(p4.add_run("，節省電費 "), size=12) 
        set_font_kai(p4.add_run(f"{(savings_money/10000):.1f} 萬元/年"), size=12, color=RGBColor(255, 0, 0)) # 紅字
        set_font_kai(p4.add_run("。"), size=12)

        p5 = doc.add_paragraph()
        set_font_kai(p5.add_run("2. 投資費用：高效率變壓器汰換投資費用預估約 "), size=12)
        set_font_kai(p5.add_run(f"{(invest_cost/10000):.1f} 萬元"), size=12, color=RGBColor(255, 0, 0)) # 紅字
        set_font_kai(p5.add_run(" (實際金額依廠商報價為主)。"), size=12, is_bold=True)

        p6 = doc.add_paragraph()
        set_font_kai(p6.add_run("3. 回收年限："), size=12)
        set_font_kai(p6.add_run(f"{(invest_cost/10000):.1f} 萬元 ÷ {(savings_money/10000):.1f} 萬元/年 = "), size=12, color=RGBColor(255, 0, 0)) # 紅字
        set_font_kai(p6.add_run(f"{payback_year:.1f} 年"), size=12, color=RGBColor(255, 0, 0)) # 紅字
        set_font_kai(p6.add_run("(註：回收年限會依報價廠家不同而有所增減)。"), size=12, is_bold=True)
output = io.BytesIO()
doc.save(output)
report_data = output.getvalue()

if 'report_warehouse' in st.session_state:
    st.session_state['report_warehouse']["1. 變壓器分析報告"] = report_data

st.success("✅ 變壓器報告已生成！您可以在左側側邊欄打包下載。")
# 提供單份報告下載按鈕 (確保這兩行也完全靠左)
st.download_button(
    label="💾 下載此份變壓器報告",
    data=report_data,
    file_name="變壓器汰換效益分析.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
