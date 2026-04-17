import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import io

# --- 1. 通用工具函數 ---
def set_font_kai(run, size=14, is_bold=False, color=RGBColor(0, 0, 0)):
    run.font.name = '標楷體'
    run.font.size = Pt(size)
    run.font.bold = is_bold
    run.font.color.rgb = color
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

# --- 2. 數據抓取邏輯 (多電號修正版) ---
def fetch_exact_data():
    info = {"comp": "未抓到名稱", "area": "0", "air_area": "0", "emp": "0", "hours": "0", "date": "115年1月1日"}
    elec_list = [] 

    if 'global_excel' in st.session_state and st.session_state['global_excel'] is not None:
        try:
            file = st.session_state['global_excel']
            xl = pd.ExcelFile(file)
            
            # --- A. 抓基本資料 (sheet_b) ---
            sheet_b = next((s for s in xl.sheet_names if "三" in s or "基本資料" in s), None)
            if sheet_b:
                df_b = pd.read_excel(file, sheet_name=sheet_b, header=None)
                # ... 這裡保留你原本抓 [emp, hours, area, air_area] 的 get_near_value 邏輯 ...

            # --- B. 遍歷所有五之二 (多電號) ---
            p_sheets = [s for s in xl.sheet_names if "五之二" in s]
            for i, s_name in enumerate(p_sheets):
                df_p = pd.read_excel(file, sheet_name=s_name, header=None)
                
                # 抓名稱 (只在第一個電號抓一次)
                if i == 0:
                    try:
                        name_val = str(df_p.iloc[5, 4]).strip()
                        if name_val != "nan":
                            info["comp"] = name_val.split('(')[0].split('（')[0]
                    except: pass

                # 建立該電號的數據字典
                e_data = {
                    "elec_id": str(df_p.iloc[5, 2]).strip(),
                    "contract_cap": "0", "total_kwh": "0", "total_fee": "0", 
                    "avg_pf": "0", "avg_price": "0", "volt": "22.8",
                    "trans_cap": "0", "cap_cap": "0", "peak_max": "0", "offpeak_max": "0"
                }

                # 抓取該分頁的電力座標 (座標依照你之前的定義)
                try:
                    e_data["contract_cap"] = str(int(float(df_p.iloc[9, 2])))
                    kwh = float(df_p.iloc[21, 11])
                    e_data["total_kwh"] = f"{int(kwh):,d}"
                    fee = float(df_p.iloc[21, 14])
                    e_data["total_fee"] = f"{int(fee):,d}"
                    if kwh > 0: e_data["avg_price"] = str(round(fee / kwh, 2))
                    e_data["avg_pf"] = str(int(float(df_p.iloc[22, 13])))
                    
                    # 需量最大值 (D10~D21 為尖峰, G10~G21 為離峰)
                    p_vals = [float(df_p.iloc[r, 3]) for r in range(9, 21) if pd.notnull(df_p.iloc[r, 3]) and str(df_p.iloc[r, 3]).strip() not in ["-", "0"]]
                    if p_vals: e_data["peak_max"] = str(int(max(p_vals)))
                    o_vals = [float(df_p.iloc[r, 6]) for r in range(9, 21) if pd.notnull(df_p.iloc[r, 6]) and str(df_p.iloc[r, 6]).strip() not in ["-", "0"]]
                    if o_vals: e_data["offpeak_max"] = str(int(max(o_vals)))
                except: pass

                # --- C. 只有第一個電號處理「表八」 ---
                if i == 0:
                    sheet_8 = next((s for s in xl.sheet_names if "八" in s), None)
                    if sheet_8:
                        df_8 = pd.read_excel(file, sheet_name=sheet_8, header=None)
                        # 變壓器加總 (第8列)
                        t_sum = sum([float(df_8.iloc[7, col]) for col in range(5, len(df_8.columns)) if pd.notnull(df_8.iloc[7, col]) and isinstance(df_8.iloc[7, col], (int,float))])
                        e_data["trans_cap"] = f"{int(t_sum):,d}"
                        # 電容器加總 (第23列)
                        c_sum = sum([float(df_8.iloc[22, col]) for col in range(5, len(df_8.columns)) if pd.notnull(df_8.iloc[22, col]) and isinstance(df_8.iloc[22, col], (int,float))])
                        e_data["cap_cap"] = str(int(c_sum))

                elec_list.append(e_data)
        except Exception as e:
            st.error(f"解析發生錯誤: {e}")
            
    return info, elec_list

    return info, elec_list # 這裡現在回傳兩個東西
            
           # --- 處理「表五之二」 ---
            sheet_p = next((s for s in xl.sheet_names if "五之二" in s), None)
            if sheet_p:
                df_p = pd.read_excel(file, sheet_name=sheet_p, header=None)
                
                # 1. 修正名稱抓取：直接抓 E6 (索引 5, 4)
                try:
                    name_val = str(df_p.iloc[5, 4]).strip()
                    if name_val != "nan" and name_val != "":
                        # 移除括號後的內容（例如公司編號）
                        info["comp"] = name_val.split('(')[0].split('（')[0]
                except:
                    pass

                # 2. 抓電號 C6 (索引 5, 2)
                # ... 後面的電號、需量、合計邏輯維持不變 ...
                # 台電電號 C6 (索引 5, 2)
                info["elec_id"] = str(df_p.iloc[5, 2]).strip()
                
                # 契約容量 C10 (索引 9, 2)
                try: info["contract_cap"] = str(int(float(df_p.iloc[9, 2])))
                except: pass

                # 年總用電度 L22 (索引 21, 11)
                try: 
                    kwh = float(df_p.iloc[21, 11])
                    info["total_kwh"] = f"{int(kwh):,d}"
                except: kwh = 0

                # 年總金額 O22 (索引 21, 14)
                try: 
                    fee = float(df_p.iloc[21, 14])
                    info["total_fee"] = f"{int(fee):,d}"
                except: fee = 0

                # 平均單價 O22 / L22
                if kwh > 0 and fee > 0:
                    info["avg_price"] = str(round(fee / kwh, 2))

                # 平均功因 N23 (索引 22, 13)
                try: info["avg_pf"] = str(int(float(df_p.iloc[22, 13])))
                except: pass

                # --- 需量精準抓取 ---
                # 尖峰最高需量 D10~D21 (索引 9~20, 欄 3)
                try:
                    peak_vals = [float(df_p.iloc[r, 3]) for r in range(9, 21) if pd.notnull(df_p.iloc[r, 3]) and str(df_p.iloc[r, 3]).strip() not in ["-", "0", "0.0"]]
                    if peak_vals: info["peak_max"] = str(int(max(peak_vals)))
                except: pass

                # 離峰最高需量 G10~G21 (索引 9~20, 欄 6)
                try:
                    off_vals = [float(df_p.iloc[r, 6]) for r in range(9, 21) if pd.notnull(df_p.iloc[r, 6]) and str(df_p.iloc[r, 6]).strip() not in ["-", "0", "0.0"]]
                    if off_vals: info["offpeak_max"] = str(int(max(off_vals)))
                except: pass

            # --- 2. 處理「表八」(變壓器總和、電容器總和) ---
            sheet_8 = next((s for s in xl.sheet_names if "八" in s), None)
            if sheet_8:
                df_8 = pd.read_excel(file, sheet_name=sheet_8, header=None)
                
                # 變壓器容量總和：第 8 列 (索引 7)，從 F 欄 (索引 5) 往後掃描
                try:
                    t_sum = 0
                    for col in range(5, len(df_8.columns)):
                        v = df_8.iloc[7, col]
                        if pd.notnull(v) and isinstance(v, (int, float)):
                            t_sum += v
                    info["trans_cap"] = f"{int(t_sum):,d}"
                except: pass

                # 電容器容量總和：第 23 列 (索引 22)，從 F 欄 (索引 5) 往後掃描
                try:
                    cap_sum = 0
                    for col in range(5, len(df_8.columns)):
                        v = df_8.iloc[22, col]
                        if pd.notnull(v) and isinstance(v, (int, float)):
                            cap_sum += v
                    info["cap_cap"] = str(int(cap_sum))
                except: pass

            # --- 處理「基本資料」(人數、面積、工時) ---
            sheet_b = next((s for s in xl.sheet_names if "三" in s or "基本資料" in s), None)
            if sheet_b:
                df_b = pd.read_excel(file, sheet_name=sheet_b, header=None)
                
                def get_near_value(items, keyword, min_val=0):
                    import re
                    for i, item in enumerate(items):
                        if keyword in str(item):
                            for target in items[i+1 : i+5]:
                                if target is None or str(target).lower() == "nan": continue
                                clean = str(target).replace(",", "").replace(" ", "")
                                matches = re.findall(r"[-+]?\d*\.\d+|\d+", clean)
                                if matches:
                                    try:
                                        num = int(round(float(matches[0])))
                                        if num > min_val: return f"{num:,d}"
                                    except: continue
                    return None

                for r in range(len(df_b)):
                    row_list = list(df_b.iloc[r, :])
                    row_str = "".join([str(i) for i in row_list])
                    if "員工人數" in row_str:
                        res = get_near_value(row_list, "員工人數")
                        if res: info["emp"] = res
                    if "全年工作時數" in row_str:
                        res = get_near_value(row_list, "全年工作時數")
                        if res: info["hours"] = res
                    if "總樓地板面積" in row_str:
                        res = get_near_value(row_list, "總樓地板面積", min_val=100)
                        if res: info["area"] = res
                    if "總空調使用面積" in row_str:
                        res = get_near_value(row_list, "總空調使用面積", min_val=100)
                        if res: info["air_area"] = res

        except Exception as e:
            st.error(f"解析發生錯誤: {e}")
            
    return info

# --- 3. 介面 ---
st.title("📋 節能診斷自動化工具")

# 只調用這一次
info_result, elec_systems = fetch_exact_data()

with st.expander("🔍 檢視/微調自動抓取資料"):
    ec1, ec2 = st.columns(2)
    with ec1:
        # 使用 info_result 內的資料
        v_comp = st.text_input("用戶名稱", info_result["comp"])
        v_area = st.text_input("總面積", info_result["area"])
        v_air = st.text_input("空調面積", info_result["air_area"])
    with ec2:
        v_emp = st.text_input("員工人數", info_result["emp"])
        v_hours = st.text_input("工作時數", info_result["hours"])

v_date = st.text_input("📅 診斷日期", info_result["date"])

# 電力系統 Tabs 保持你寫的那樣，但要補齊所有欄位 (單價, 需量等)
# ... (這裡維持你截圖中 col1, col2, col3 的邏輯) ...

st.markdown("### ⚡ 電力系統設備資料")
# 根據電號數量產生標籤頁
if elec_systems:
    tabs = st.tabs([f"電號 {e['elec_id']}" for e in elec_systems])
    for i, tab in enumerate(tabs):
        with tab:
            col1, col2, col3 = st.columns(3)
            # 注意：這裡的變數要存回 elec_systems[i] 裡面，且 key 要唯一
            elec_systems[i]['elec_id'] = col1.text_input("台電電號", elec_systems[i]['elec_id'], key=f"id_{i}")
            elec_systems[i]['total_kwh'] = col1.text_input("年總用電度", elec_systems[i]['total_kwh'], key=f"kwh_{i}")
            elec_systems[i]['contract_cap'] = col2.text_input("契約容量", elec_systems[i]['contract_cap'], key=f"cap_{i}")
            elec_systems[i]['trans_cap'] = col2.text_input("主變壓器容量", elec_systems[i]['trans_cap'], key=f"trans_{i}")
            # ... 依此類推補完其他欄位 ...

# --- 4. 封裝 Word 生成邏輯 ---
def generate_docx(comp, area, air, emp, hours, date, elecs):
    doc = Document()
    p_t1 = doc.add_paragraph(); set_font_kai(p_t1.add_run("二、能源用戶概述"), is_bold=True)
    p_t2 = doc.add_paragraph(); set_font_kai(p_t2.add_run("  2-1. 用戶簡介"), is_bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(28)
    # 注意這裡：要用傳進來的變數名稱
    set_font_kai(p.add_run(comp), color=RGBColor(255, 0, 0))
    set_font_kai(p.add_run("總建物面積"))
    set_font_kai(p.add_run(area), color=RGBColor(255, 0, 0))
    set_font_kai(p.add_run("平方公尺，空調使用面積"))
    set_font_kai(p.add_run(air), color=RGBColor(255, 0, 0))
    set_font_kai(p.add_run("平方公尺，能源使用主要以"))
    set_font_kai(p.add_run("電力"), color=RGBColor(255, 0, 0))
    set_font_kai(p.add_run("為主，員工約有"))
    set_font_kai(p.add_run(emp), color=RGBColor(255, 0, 0))
    set_font_kai(p.add_run("人，全年使用時間約"))
    set_font_kai(p.add_run(hours), color=RGBColor(255, 0, 0))
    set_font_kai(p.add_run("小時，"))
    set_font_kai(p.add_run(date), color=RGBColor(255, 0, 0)) 
    set_font_kai(p.add_run("經由實地查訪貴單位之公用系統使用情形及輔導診斷概述如下："))

    # --- 關鍵：循環必須縮進在函數裡面 ---
    for i, e in enumerate(elecs):
        doc.add_paragraph() # 空行
        set_font_kai(doc.add_paragraph().add_run(f"1.{i+1} 電力系統 (電號：{e['elec_id']})："), is_bold=True)

        table = doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'
        
        # 合併第一列並填入電號
        cell_id = table.cell(0, 0); cell_id.merge(table.cell(0, 2))
        p_id = cell_id.paragraphs[0]
        set_font_kai(p_id.add_run("台電電號："), size=12)
        set_font_kai(p_id.add_run(e['elec_id']), size=12, color=RGBColor(255, 0, 0))

        # 這裡的變數全部都要改成 e['欄位名稱']，否則會當機
        r1 = table.rows[1].cells
        set_font_kai(r1[0].paragraphs[0].add_run(f"契約型式：高壓 3 段式"), size=12)
        set_font_kai(r1[1].paragraphs[0].add_run(f"契約容量：{e['contract_cap']} [kW]"), size=12)
        set_font_kai(r1[2].paragraphs[0].add_run(f"台電供電電壓：{e.get('volt', '22.8')} [kV]"), size=12)

        r2 = table.rows[2].cells
        set_font_kai(r2[0].paragraphs[0].add_run(f"主變壓器總裝置容量：{e['trans_cap']} [kVA]"), size=12)
        set_font_kai(r2[1].paragraphs[0].add_run(f"電容器裝置容量：{e['cap_cap']} [kVAR]"), size=12)
        set_font_kai(r2[2].paragraphs[0].add_run(f"低壓側電壓：380/220 [V]"), size=12)

        r3 = table.rows[3].cells
        set_font_kai(r3[0].paragraphs[0].add_run(f"年總用電度：{e['total_kwh']} [kWh]"), size=12)
        set_font_kai(r3[1].paragraphs[0].add_run(f"年總金額：{e['total_fee']} [元]"), size=12)
        set_font_kai(r3[2].paragraphs[0].add_run(f"平均單價：{e.get('avg_price', '0')} [元/kWh]"), size=12)

        r4 = table.rows[4].cells
        set_font_kai(r4[0].paragraphs[0].add_run(f"平均功因：{e['avg_pf']} [%]"), size=12)
        set_font_kai(r4[1].paragraphs[0].add_run(f"尖峰最高需量：{e.get('peak_max', '0')} [kW]"), size=12)
        set_font_kai(r4[2].paragraphs[0].add_run(f"離峰最高需量：{e.get('offpeak_max', '0')} [kW]"), size=12)

        for row in table.rows:
            for cell in row.cells: cell.vertical_alignment = 1

    target_stream = io.BytesIO()
    doc.save(target_stream)
    return target_stream.getvalue()

# --- 5. 下載按鈕 ---
st.markdown("---")
# 點擊時才執行生成邏輯，並傳入網頁上的最新變數
if st.download_button(
    label="💾 生成並下載用戶簡介 Word",
    data=generate_docx(v_comp, v_area, v_air, v_emp, v_hours, v_date, elec_systems),
    file_name=f"能源用戶簡介_{v_comp}.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
):
    st.success("檔案下載準備就緒！")
